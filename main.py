# import PyPDF2,docx2pdf,fitz
import PyPDF2,docx2pdf
# from aspose.pdf.text import TextAbsorber
# import docx2pdf as docx2pdf
from docx import Document
from pdf2docx import Converter
import PyPDF2
# import textwrap

pdf_file_path="C:\\Users\\RAVI\\Documents\\MSD_WORD.pdf"     # To Define The PDF File Path Here
doc_file_path="C:\\Users\\RAVI\\Documents\\MSD_WORD.docx"   # To Define The DOCS File Path Here
# # TO GET THE NO OF PAGES IN BOTH PDF AND DOCS FiLE

def No_Pages_in_Files_and_Paragraph(file_path):
    try:
        extension=str(file_path).split(".")[-1]
        pdf_page_count = 0
        doc_page_count = 0

        if extension=='pdf' or extension=='PDF':
            with open(file_path,'rb') as file_data:
                load_the_pdf_file=PyPDF2.PdfFileReader(file_data)
                no_page=load_the_pdf_file.getNumPages()
                pdf_page_count=no_page
        if extension=='docx' or extension=='DOCX':
            doc_data=Document(doc_file_path)
            print(doc_data)
            pdf_file_path = file_path.replace('.docx', '.pdf')
            docx2pdf.convert(file_path, pdf_file_path)

            # Get number of pages from the PDF
            with open(pdf_file_path, 'rb') as file:
                reader = PyPDF2.PdfFileReader(file)
                num_pages = reader.numPages
                doc_page_count=num_pages





        return {"pdf_page_count":pdf_page_count,"doc_page_count":doc_page_count}
    except Exception as error:
        return {"error": str(error), "err_line": str(error.__traceback__.tb_lineno)}

pdf_func=No_Pages_in_Files_and_Paragraph(pdf_file_path)
# doc_func=No_Pages_in_Files_and_Paragraph(doc_file_path)
print(pdf_func)
# print(doc_func)





def read_docx_words_with_attributes(docx_file_path):
    try:
        doc = Document(docx_file_path)
        words_with_attributes = []
        page_strt=0
        for paragraph in doc.paragraphs:##Paragraph Indicate The Block Of Text In Word Separated By New Line
            for run in paragraph.runs:## runs Indicates The Collections Of Runs , also Provide The Contigeous Range Of Text with in paragraph has same Character Formatting , character also include the font and other attributes
                page_strt=page_strt+1
                text = run.text
                for word in text.split():#plit The Text Into Words
                    bold = run.bold
                    italic = run.italic
                    underline = run.underline
                    font_size = run.font.size.pt if run.font.size else None
                    font_name = run.font.name
                    words_with_attributes.append({"text": word, "bold": bold, "italic": italic, "underline": underline,
                                                  "font_size": font_size, "font_name": font_name,"page_line":page_strt})


        return words_with_attributes
    except Exception as error:
        return {"error": str(error), "err_line": str(error.__traceback__.tb_lineno)}
# Example usage
docx_words_with_attributes_D = read_docx_words_with_attributes(doc_file_path)

if isinstance(docx_words_with_attributes_D,dict):
    print(docx_words_with_attributes_D)
    print("INVALID")

def convert_pdf_to_docx(pdf_file_path, docx_file_path):
    try:
        # Create a PDF to DOCX converter object
        cv = Converter(pdf_file_path)

        # Perform the conversion
        cv.convert(docx_file_path, start=0, end=None)

        # Close the converter
        cv.close()
    except Exception as error:
        return {"error": str(error), "err_line": str(error.__traceback__.tb_lineno)}

# Example usage

docx_file_path = 'C:\\Users\\RAVI\\Documents\\output1.docx'
dcon=convert_pdf_to_docx(pdf_file_path, docx_file_path)

if isinstance(dcon,dict):
    print(dcon)
    print("INVALID")
docx_words_with_attributes_P = read_docx_words_with_attributes(docx_file_path)
if isinstance(docx_words_with_attributes_P,dict):
    print(docx_words_with_attributes_P)
    print("INVALID")
    
print(docx_words_with_attributes_D)
print(docx_words_with_attributes_P)

disc_data=[]
print("Len_Docx="+str(len(docx_words_with_attributes_D)))
print("Len_PDF="+str(len(docx_words_with_attributes_P)))

for i in docx_words_with_attributes_P:
    for j in docx_words_with_attributes_D:
        if i.get('text')==j.get('text'):
            if ((i.get('bold')==None or i.get('bold')==False) and j.get('bold')==True) or (i.get('bold')==True and (j.get('bold')==False )):
                disc_data.append({"text": i.get('text'), 'docx_bold': i.get('bold'), "pdf_bold": j.get("bold")})
            if  ((i.get('italic')==None or i.get('italic')==False) and j.get('italic')==True) or (i.get('italic')==True and (j.get('italic')==False )):#underline
                disc_data.append({"text": i.get('text'), 'docx_italic': i.get('italic'), "pdf_italic": j.get("italic")})

            if ((i.get('underline')==None or i.get('underline')==False) and j.get('underline')==True) or (i.get('underline')==True and (j.get('underline')==False or j.get('underline')==None)):
                disc_data.append( {"text": i.get('text'), 'docx_underline': i.get('underline'), "pdf_underline": j.get("underline")})
print(disc_data)
print(len(disc_data))